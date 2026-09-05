from copy import deepcopy
from pathlib import Path

from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from app import word_renderer


EMU_PER_PT = 12700


def _safe_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _page_relative_image(doc, image_info, element, page_rect):
    """Render one Gemini image element at its page coordinates.

    The production PDF-to-Word path is app.main -> word_renderer.render_editable_pdf.
    Earlier fixes changed app/renderer.py, which that endpoint never imports. This
    adapter patches the actual word_renderer image hook and uses Gemini's page-space
    x/y/w/h for the displayed geometry instead of the extracted image's intrinsic
    PDF placement rectangle.
    """
    path = image_info["path"]
    page_w = page_rect.width
    page_h = page_rect.height

    # Gemini layout values are normalized to the current page. If dimensions are
    # missing, fall back to the extracted PDF image rectangle.
    has_geometry = all(k in element for k in ("x", "y", "w", "h"))
    if has_geometry:
        x_pt = _safe_float(element.get("x")) * page_w
        y_pt = _safe_float(element.get("y")) * page_h
        width_pt = max(0.01, _safe_float(element.get("w")) * page_w)
        height_pt = max(0.01, _safe_float(element.get("h")) * page_h)
    else:
        rect = image_info["rect"]
        x_pt, y_pt = rect.x0, rect.y0
        width_pt, height_pt = rect.width, rect.height

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    run = p.add_run()
    inline = run.add_picture(
        str(path),
        width=Inches(width_pt / 72.0),
        height=Inches(height_pt / 72.0),
    )
    drawing = inline._inline

    anchor = OxmlElement("wp:anchor")
    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251658240",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "0",
        "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    off_h = OxmlElement("wp:posOffset")
    # Do not clamp the coordinate: negative PDF/Gemini coordinates are meaningful
    # when an image intentionally crosses a page edge.
    off_h.text = str(round(x_pt * EMU_PER_PT))
    pos_h.append(off_h)
    anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    off_v = OxmlElement("wp:posOffset")
    off_v.text = str(round(y_pt * EMU_PER_PT))
    pos_v.append(off_v)
    anchor.append(pos_v)

    for child in list(drawing):
        anchor.append(deepcopy(child))
    anchor.append(OxmlElement("wp:wrapNone"))
    drawing.getparent().replace(drawing, anchor)


def _patched_add_floating_image(doc, image_info, element, page_rect):
    return _page_relative_image(doc, image_info, element, page_rect)


# word_renderer.render_editable_pdf resolves _add_floating_image at call time,
# so patching this single hook changes the actual production endpoint without
# duplicating the rest of the renderer.
word_renderer._add_floating_image = _patched_add_floating_image
render_editable_pdf = word_renderer.render_editable_pdf
