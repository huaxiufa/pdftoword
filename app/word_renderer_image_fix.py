from pathlib import Path
from copy import deepcopy
from PIL import Image
import fitz
from docx.oxml import OxmlElement
from app import word_renderer as _wr


def _extract_images(page, work_dir):
    result = []
    seen = set()
    for info in page.get_images(full=True):
        try:
            xref = info[0]
            smask = info[1] if len(info) > 1 else 0
            if smask:
                base = fitz.Pixmap(page.parent, xref)
                mask = fitz.Pixmap(page.parent, smask)
                pix = fitz.Pixmap(base, mask)
                path = Path(work_dir) / f"p{page.number + 1}_img{xref}.png"
                pix.save(str(path))
            else:
                data = page.parent.extract_image(xref)
                path = Path(work_dir) / f"p{page.number + 1}_img{xref}.{data.get('ext', 'png')}"
                path.write_bytes(data['image'])
            for rect in page.get_image_rects(xref):
                clipped = rect & page.rect
                if clipped.is_empty:
                    continue
                if clipped != rect:
                    src = Image.open(path).convert('RGBA')
                    sx = src.width / rect.width
                    sy = src.height / rect.height
                    box = (
                        max(0, round((clipped.x0 - rect.x0) * sx)),
                        max(0, round((clipped.y0 - rect.y0) * sy)),
                        min(src.width, round((clipped.x1 - rect.x0) * sx)),
                        min(src.height, round((clipped.y1 - rect.y0) * sy)),
                    )
                    cropped = Path(work_dir) / f"p{page.number + 1}_img{xref}_{len(result)}.png"
                    src.crop(box).save(cropped, 'PNG')
                    path = cropped
                    rect = clipped
                key = (xref, round(rect.x0, 3), round(rect.y0, 3), round(rect.x1, 3), round(rect.y1, 3))
                if key not in seen:
                    result.append({'path': path, 'rect': rect, 'xref': xref})
                    seen.add(key)
        except Exception:
            continue
    result.sort(key=lambda x: (x['rect'].y0, x['rect'].x0))
    return result


def _add_floating_image(cell, image_info, element, page_rect):
    actual = image_info['rect']
    path = image_info['path']
    width_pt = max(1.0, actual.width)
    height_pt = max(1.0, actual.height)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    run = p.add_run()
    inline = run.add_picture(str(path), width=_wr.Inches(width_pt / 72), height=_wr.Inches(height_pt / 72))
    drawing = inline._inline

    anchor = OxmlElement('wp:anchor')
    for key, value in {
        'distT': '0', 'distB': '0', 'distL': '0', 'distR': '0',
        'simplePos': '0', 'relativeHeight': '251658240',
        'behindDoc': '0', 'locked': '0', 'layoutInCell': '0',
        'allowOverlap': '1'
    }.items():
        anchor.set(key, value)

    simple = OxmlElement('wp:simplePos')
    simple.set('x', '0')
    simple.set('y', '0')
    anchor.append(simple)

    pos_h = OxmlElement('wp:positionH')
    pos_h.set('relativeFrom', 'page')
    offset_h = OxmlElement('wp:posOffset')
    offset_h.text = str(round(actual.x0 * 12700))
    pos_h.append(offset_h)
    anchor.append(pos_h)

    pos_v = OxmlElement('wp:positionV')
    pos_v.set('relativeFrom', 'page')
    offset_v = OxmlElement('wp:posOffset')
    offset_v.text = str(round(actual.y0 * 12700))
    pos_v.append(offset_v)
    anchor.append(pos_v)

    for child in list(drawing):
        anchor.append(deepcopy(child))
    anchor.append(OxmlElement('wp:wrapNone'))
    drawing.getparent().replace(drawing, anchor)


_wr._extract_images = _extract_images
_wr._add_floating_image = _add_floating_image
render_editable_pdf = _wr.render_editable_pdf
