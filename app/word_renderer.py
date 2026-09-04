from pathlib import Path
from copy import deepcopy

import fitz
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _safe_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _safe_int(value, default=0):
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return default


def _color(value, default="000000"):
    if not isinstance(value, str):
        return default
    value = value.strip().lstrip("#")
    return value.upper() if len(value) == 6 else default


def _page_elements(gem_page):
    if not isinstance(gem_page, dict):
        return []
    elements = []
    direct = gem_page.get("elements")
    if isinstance(direct, list):
        elements.extend(x for x in direct if isinstance(x, dict))
    columns = gem_page.get("columns")
    if isinstance(columns, list):
        for col in columns:
            if isinstance(col, dict) and isinstance(col.get("elements"), list):
                elements.extend(x for x in col["elements"] if isinstance(x, dict))
    return sorted(elements, key=lambda e: (_safe_float(e.get("y")), _safe_float(e.get("x"))))


def _column_elements(gem_page):
    columns = gem_page.get("columns") if isinstance(gem_page, dict) else None
    if not isinstance(columns, list) or not columns:
        return [{"x": 0.0, "y": 0.0, "w": 1.0, "h": 1.0, "elements": _page_elements(gem_page)}]
    result = []
    for col in columns:
        if not isinstance(col, dict):
            continue
        items = col.get("elements", [])
        if isinstance(items, list):
            result.append({
                "x": _safe_float(col.get("x")),
                "y": _safe_float(col.get("y")),
                "w": max(0.01, _safe_float(col.get("w"), 1.0)),
                "h": max(0.01, _safe_float(col.get("h"), 1.0)),
                "elements": sorted((x for x in items if isinstance(x, dict)), key=lambda e: (_safe_float(e.get("y")), _safe_float(e.get("x")))),
            })
    return result or [{"x": 0.0, "y": 0.0, "w": 1.0, "h": 1.0, "elements": _page_elements(gem_page)}]


def _set_cell_text(cell, text, element=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run("" if text is None else str(text))
    if element:
        run.bold = bool(element.get("bold"))
        run.italic = bool(element.get("italic"))
        run.underline = bool(element.get("underline"))
        run.font.size = Pt(max(6, _safe_float(element.get("font_size"), 9)))
        try:
            run.font.name = str(element.get("font_family") or "Arial")
            run.font.color.rgb = RGBColor.from_string(_color(element.get("color")))
        except Exception:
            pass


def _add_text(cell, element):
    text = str(element.get("text") or "")
    if not text.strip():
        return
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1
    align = str(element.get("align") or "left").lower()
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bool(element.get("bold"))
    run.italic = bool(element.get("italic"))
    run.underline = bool(element.get("underline"))
    run.font.size = Pt(max(6, _safe_float(element.get("font_size"), 10.5)))
    try:
        run.font.name = str(element.get("font_family") or "Arial")
        run.font.color.rgb = RGBColor.from_string(_color(element.get("color")))
    except Exception:
        pass


def _normalize_table(element):
    cells = element.get("cells")
    normalized = []
    if isinstance(cells, list):
        for item in cells:
            if not isinstance(item, dict):
                continue
            normalized.append({
                "row": max(0, _safe_int(item.get("row"))),
                "col": max(0, _safe_int(item.get("col"))),
                "row_span": max(1, _safe_int(item.get("row_span"), 1)),
                "col_span": max(1, _safe_int(item.get("col_span"), 1)),
                "text": "" if item.get("text") is None else str(item.get("text")),
            })
    if normalized:
        max_row = max(x["row"] + x["row_span"] for x in normalized)
        max_col = max(x["col"] + x["col_span"] for x in normalized)
        return max_row, max_col, normalized

    rows = element.get("rows")
    if not isinstance(rows, list):
        return 0, 0, []
    clean = []
    for row in rows:
        if isinstance(row, list):
            clean.append(row)
        elif isinstance(row, dict):
            clean.append([row.get("text", "")])
        else:
            clean.append([row])
    cols = max((len(r) for r in clean), default=0)
    return len(clean), cols, [
        {"row": r, "col": c, "row_span": 1, "col_span": 1, "text": row[c] if c < len(row) else ""}
        for r, row in enumerate(clean) for c in range(cols)
    ]


def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "nil")


def _set_table_layout(table, width_pt):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table.width = Inches(max(1.0, width_pt) / 72.0)


def _add_table(cell, element):
    rows, cols, source = _normalize_table(element)
    if not rows or not cols:
        return
    table = cell.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    table_width_pt = max(72.0, _safe_float(element.get("w"), 1.0) * 600.0)
    col_width = table_width_pt / cols
    for r in range(rows):
        row = table.rows[r]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for c in range(cols):
            target = table.cell(r, c)
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            target.width = Inches(col_width / 72.0)

    for item in sorted(source, key=lambda x: (x["row"], x["col"])):
        if item["row"] < rows and item["col"] < cols:
            _set_cell_text(table.cell(item["row"], item["col"]), item["text"], element)

    occupied = set()
    for item in sorted(source, key=lambda x: (x["row"], x["col"], -x["row_span"], -x["col_span"])):
        r, c = item["row"], item["col"]
        rs, cs = item["row_span"], item["col_span"]
        cells = {(rr, cc) for rr in range(r, min(rows, r + rs)) for cc in range(c, min(cols, c + cs))}
        if rs == 1 and cs == 1:
            occupied.add((r, c))
            continue
        if cells & (occupied - {(r, c)}):
            continue
        try:
            table.cell(r, c).merge(table.cell(min(rows - 1, r + rs - 1), min(cols - 1, c + cs - 1)))
            occupied.update(cells)
        except Exception:
            pass
    cell.add_paragraph().paragraph_format.space_after = Pt(0)


def _extract_images(page, work_dir):
    result = []
    seen = set()
    for info in page.get_images(full=True):
        try:
            xref = info[0]
            smask = info[1] if len(info) > 1 else 0
            if smask:
                base = fitz.Pixmap(page.parent, xref)
                mask = fitz.Pixmap(page.parent, smask)
                pix = fitz.Pixmap(base, mask)
                path = Path(work_dir) / f"p{page.number + 1}_img{xref}.png"
                pix.save(str(path))
            else:
                data = page.parent.extract_image(xref)
                ext = data.get("ext", "png")
                path = Path(work_dir) / f"p{page.number + 1}_img{xref}.{ext}"
                path.write_bytes(data["image"])
            for rect in page.get_image_rects(xref):
                # Word/LibreOffice may shift an anchored image that extends outside
                # the page. Crop only the out-of-page part so the remaining visible
                # pixels can be anchored exactly at the PDF page boundary.
                clipped = rect & page.rect
                if clipped.is_empty:
                    continue
                if clipped != rect:
                    src = Image.open(path).convert("RGBA")
                    sx = src.width / rect.width
                    sy = src.height / rect.height
                    box = (
                        max(0, round((clipped.x0 - rect.x0) * sx)),
                        max(0, round((clipped.y0 - rect.y0) * sy)),
                        min(src.width, round((clipped.x1 - rect.x0) * sx)),
                        min(src.height, round((clipped.y1 - rect.y0) * sy)),
                    )
                    cropped_path = Path(work_dir) / f"p{page.number + 1}_img{xref}_{len(result)}.png"
                    src.crop(box).save(cropped_path, "PNG")
                    path = cropped_path
                    rect = clipped
                key = (xref, round(rect.x0, 3), round(rect.y0, 3), round(rect.x1, 3), round(rect.y1, 3))
                if key not in seen:
                    result.append({"path": path, "rect": rect, "xref": xref})
                    seen.add(key)
        except Exception:
            continue
    result.sort(key=lambda item: (item["rect"].y0, item["rect"].x0))
    return result


def _match_image(image_map, element, page_rect):
    if not image_map:
        return None
    idx = _safe_int(element.get("image_index"), -1)
    if 0 <= idx < len(image_map):
        return image_map[idx]
    ex = _safe_float(element.get("x"), 0.0) * page_rect.width
    ey = _safe_float(element.get("y"), 0.0) * page_rect.height
    ew = _safe_float(element.get("w"), 0.0) * page_rect.width
    eh = _safe_float(element.get("h"), 0.0) * page_rect.height
    ecx, ecy = ex + ew / 2, ey + eh / 2
    def score(item):
        r = item["rect"]
        cx, cy = (r.x0 + r.x1) / 2, (r.y0 + r.y1) / 2
        return ((cx - ecx) ** 2 + (cy - ecy) ** 2) ** 0.5
    return min(image_map, key=score)


def _add_floating_image(cell, image_info, element, page_rect):
    """Insert an absolutely positioned image into the current page's layout cell."""
    path = image_info["path"]
    actual = image_info["rect"]
    width_pt = max(1.0, actual.width)
    height_pt = max(1.0, actual.height)
    x_pt, y_pt = actual.x0, actual.y0
    if actual.width <= 1 or actual.height <= 1:
        width_pt = max(1.0, _safe_float(element.get("w"), 0.1) * page_rect.width)
        height_pt = max(1.0, _safe_float(element.get("h"), 0.1) * page_rect.height)
        x_pt = _safe_float(element.get("x")) * page_rect.width
        y_pt = _safe_float(element.get("y")) * page_rect.height

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width_pt / 72.0), height=Inches(height_pt / 72.0))
    drawing = inline._inline

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251658240")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "0")
    anchor.set("allowOverlap", "1")

    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    off_h = OxmlElement("wp:posOffset")
    off_h.text = str(round(x_pt * 12700))
    pos_h.append(off_h)
    anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    off_v = OxmlElement("wp:posOffset")
    off_v.text = str(round(y_pt * 12700))
    pos_v.append(off_v)
    anchor.append(pos_v)

    for child in list(drawing):
        anchor.append(deepcopy(child))

    anchor.append(OxmlElement("wp:wrapNone"))
    drawing.getparent().replace(drawing, anchor)


def _render_column(cell, elements, image_map, page_rect):
    cell.text = ""
    for element in elements:
        kind = str(element.get("type") or "text").lower()
        if kind in {"text", "heading", "paragraph", "bullet"}:
            item = dict(element)
            if kind == "bullet":
                item["text"] = "• " + str(item.get("text") or "")
            _add_text(cell, item)
        elif kind == "table":
            _add_table(cell, element)
        elif kind == "image":
            info = _match_image(image_map, element, page_rect)
            if info:
                _add_floating_image(cell, info, element, page_rect)
        elif kind == "line":
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("────────────────────────────────")
            run.font.size = Pt(6)


def _remove_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")


def render_editable_pdf(source_pdf, layout, output):
    source_pdf = Path(source_pdf)
    output = Path(output)
    pages = layout.get("pages", []) if isinstance(layout, dict) else []
    doc = Document()
    first = True
    with fitz.open(source_pdf) as pdf:
        for page_index, page in enumerate(pdf):
            if not first:
                doc.add_section(WD_SECTION.NEW_PAGE)
            first = False
            section = doc.sections[-1]
            section.page_width = Inches(page.rect.width / 72.0)
            section.page_height = Inches(page.rect.height / 72.0)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)

            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.allow_autofit = False
            table.width = Inches(page.rect.width / 72.0)
            cell = table.cell(0, 0)
            cell.width = Inches(page.rect.width / 72.0)
            _remove_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            page_layout = pages[page_index] if page_index < len(pages) and isinstance(pages[page_index], dict) else {}
            elements = _page_elements(page_layout)
            work_dir = output.parent / f".pdf_images_{page_index + 1}"
            work_dir.mkdir(parents=True, exist_ok=True)
            image_map = _extract_images(page, work_dir)
            _render_column(cell, elements, image_map, page.rect)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
