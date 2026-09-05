from pathlib import Path
from copy import deepcopy

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

EMU_PER_PT = 12700


def _float(v, default=0.0):
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def _hex(v):
    s = str(v or "000000").lstrip("#")
    return s.upper() if len(s) == 6 else "000000"


def _anchor_picture(doc, path, x, y, w, h):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    inline = r.add_picture(str(path), width=Inches(max(0.01, w) / 72), height=Inches(max(0.01, h) / 72))
    drawing = inline._inline
    anchor = OxmlElement("wp:anchor")
    for k, v in {"distT":"0","distB":"0","distL":"0","distR":"0","simplePos":"0","relativeHeight":"251658240","behindDoc":"0","locked":"0","layoutInCell":"0","allowOverlap":"1"}.items():
        anchor.set(k, v)
    simple = OxmlElement("wp:simplePos"); simple.set("x", "0"); simple.set("y", "0"); anchor.append(simple)
    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom", "page")
    oh = OxmlElement("wp:posOffset"); oh.text = str(int(x * EMU_PER_PT)); ph.append(oh); anchor.append(ph)
    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom", "page")
    ov = OxmlElement("wp:posOffset"); ov.text = str(int(y * EMU_PER_PT)); pv.append(ov); anchor.append(pv)
    for child in list(drawing): anchor.append(deepcopy(child))
    anchor.append(OxmlElement("wp:wrapNone"))
    drawing.getparent().replace(drawing, anchor)


def _text_box(doc, text, bbox, spans, page_w):
    x0, y0, x1, y1 = bbox
    w, h = max(1, x1-x0), max(1, y1-y0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(max(0, x0))
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    # Native PDF text is emitted in reading order, but its exact vertical
    # location is retained with an explicit top spacer before the paragraph.
    previous = getattr(doc, "_pdf_last_y", 0.0)
    gap = max(0.0, y0 - previous)
    p.paragraph_format.space_before = Pt(gap)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if spans:
        for sp in spans:
            run = p.add_run(str(sp.get("text") or ""))
            size = max(5.0, _float(sp.get("size"), 10.0))
            run.font.size = Pt(size)
            font = str(sp.get("font") or "Arial")
            run.font.name = font
            flags = int(sp.get("flags", 0))
            run.bold = bool(flags & 16)
            run.italic = bool(flags & 2)
            try: run.font.color.rgb = RGBColor.from_string(_hex(sp.get("color")))
            except Exception: pass
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    doc._pdf_last_y = max(previous, y1)


def _extract_image(page, xref, out_dir, index):
    data = page.parent.extract_image(xref)
    ext = data.get("ext", "png")
    path = Path(out_dir) / f"p{page.number+1}_img{index}.{ext}"
    # Preserve PDF transparency. extract_image() alone drops the soft mask.
    mask = None
    try:
        info = page.parent.get_xref_object(xref, compressed=False)
        sm = None
        import re
        m = re.search(r"/SMask\s+(\d+)\s+0\s+R", info)
        if m: sm = int(m.group(1))
        if sm:
            base = fitz.Pixmap(page.parent, xref)
            mpx = fitz.Pixmap(page.parent, sm)
            pix = fitz.Pixmap(base, mpx)
            pix.save(str(path))
            return path
    except Exception:
        pass
    path.write_bytes(data["image"])
    return path


def render_editable_pdf(source_pdf, layout, output):
    pdf = fitz.open(source_pdf)
    doc = Document()
    work = Path(output).parent / f".native-pdf-images-{Path(output).stem}"
    work.mkdir(parents=True, exist_ok=True)
    try:
        for page_index, page in enumerate(pdf):
            if page_index:
                doc.add_section(WD_SECTION.NEW_PAGE)
            sec = doc.sections[-1]
            sec.page_width = Pt(page.rect.width); sec.page_height = Pt(page.rect.height)
            sec.top_margin = Pt(0); sec.bottom_margin = Pt(0); sec.left_margin = Pt(0); sec.right_margin = Pt(0)
            sec.header_distance = Pt(0); sec.footer_distance = Pt(0)
            doc._pdf_last_y = 0.0

            # Use the PDF itself as the geometry authority. Gemini coordinates
            # are semantic hints only; this removes AI-dependent size drift.
            images = []
            seen = set()
            for i, info in enumerate(page.get_images(full=True)):
                xref = info[0]
                for rect in page.get_image_rects(xref):
                    key = (xref, round(rect.x0,2), round(rect.y0,2), round(rect.x1,2), round(rect.y1,2))
                    if key in seen: continue
                    seen.add(key)
                    images.append((rect, _extract_image(page, xref, work, i)))

            # Native PDF text blocks preserve real x/y/font metrics. This is
            # substantially closer to the source than reconstructing all text
            # through flowing Word table cells.
            blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT).get("blocks", [])
            items = []
            for b in blocks:
                if b.get("type") != 0: continue
                spans = [s for line in b.get("lines", []) for s in line.get("spans", [])]
                if not spans: continue
                text = "".join(str(s.get("text") or "") for s in spans)
                if text.strip(): items.append((b["bbox"], text, spans))
            items.sort(key=lambda z: (z[0][1], z[0][0]))
            for bbox, text, spans in items:
                _text_box(doc, text, bbox, spans, page.rect.width)

            for rect, path in sorted(images, key=lambda z: (z[0].y0, z[0].x0)):
                _anchor_picture(doc, path, rect.x0, rect.y0, rect.width, rect.height)

        doc.save(output)
    finally:
        pdf.close()
        for p in work.glob("*"):
            try: p.unlink()
            except Exception: pass
        try: work.rmdir()
        except Exception: pass
