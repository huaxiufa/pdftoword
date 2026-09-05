from copy import deepcopy
from io import BytesIO
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.model import DocumentModel


EMU_PER_PT = 12700


def _style_run(run, span, font_scale=1.0):
    font_name = span.font or "Arial"
    run.font.name = font_name
    run.font.size = Pt((span.size or 10) * font_scale)
    run.bold = span.bold
    run.italic = span.italic
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is not None:
        rfonts.set(qn("w:eastAsia"), font_name)


def _clean_bullet(text):
    return re.sub(r"^(?:•|·|▪|‣|◦|●|-|–|—)\s*", "", text)


def _configure_section(section, page):
    section.page_width = Inches(page.width / 72)
    section.page_height = Inches(page.height / 72)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


def _new_paragraph(doc, block, cursor_y, vertical_scale=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(max(0, block.bbox.x0))
    gap = max(0, block.bbox.y0 - cursor_y) * vertical_scale
    if gap:
        p.paragraph_format.space_before = Pt(gap)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    return p


def _add_page_anchored_image(doc, image, bbox):
    """Place an image in PDF page coordinates, outside normal text flow."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    run = p.add_run()
    inline = run.add_picture(
        BytesIO(image),
        width=Inches(bbox.width / 72),
        height=Inches(bbox.height / 72),
    )
    drawing = inline._inline

    anchor = OxmlElement("wp:anchor")
    for name, value in {
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0", "relativeHeight": "251658240",
        "behindDoc": "0", "locked": "0", "layoutInCell": "0",
        "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    offset_h = OxmlElement("wp:posOffset")
    offset_h.text = str(round(bbox.x0 * EMU_PER_PT))
    pos_h.append(offset_h)
    anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    offset_v = OxmlElement("wp:posOffset")
    offset_v.text = str(round(bbox.y0 * EMU_PER_PT))
    pos_v.append(offset_v)
    anchor.append(pos_v)

    for child in list(drawing):
        anchor.append(deepcopy(child))
    anchor.append(OxmlElement("wp:wrapNone"))
    drawing.getparent().replace(drawing, anchor)


def render_docx(model: DocumentModel, output, font_scale=1.0, vertical_scale=1.0):
    """Render editable DOCX while preserving PDF page geometry for images."""
    doc = Document()

    for page_index, page in enumerate(model.pages):
        if page_index:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        else:
            section = doc.sections[0]
        _configure_section(section, page)
        cursor_y = 0.0
        page_images = []

        # Images are not inserted inline. They are emitted after the page's
        # editable text as direct-body, page-relative floating objects.
        for block in page.blocks:
            if block.kind == "image" and block.image:
                page_images.append(block)
                continue

            p = _new_paragraph(doc, block, cursor_y, vertical_scale)

            if block.kind == "heading":
                for line in block.lines:
                    for span in line.spans:
                        r = p.add_run(span.text)
                        _style_run(r, span, font_scale)
                        r.bold = True
                cursor_y = max(cursor_y, block.bbox.y1)
                continue

            if block.kind == "list":
                p.add_run("• ")
                for line in block.lines:
                    for span in line.spans:
                        r = p.add_run(_clean_bullet(span.text))
                        _style_run(r, span, font_scale)
                cursor_y = max(cursor_y, block.bbox.y1)
                continue

            for i, line in enumerate(block.lines):
                if i:
                    p.add_run().add_break()
                for span in line.spans:
                    r = p.add_run(span.text)
                    _style_run(r, span, font_scale)
            cursor_y = max(cursor_y, block.bbox.y1)

        # Keep the image anchor paragraph in w:body, never in a table/cell.
        # relativeFrom="page" therefore uses the current section's page.
        for block in page_images:
            _add_page_anchored_image(doc, block.image, block.bbox)

    doc.save(output)
