import io
import json
import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from openpyxl import Workbook
from PIL import Image

VML_NS = "urn:schemas-microsoft-com:vml"


def merge_pdfs(paths, output):
    out = fitz.open()
    for p in paths:
        src = fitz.open(p)
        out.insert_pdf(src)
        src.close()
    out.save(output)
    out.close()


def split_pdf(path, output_dir):
    src = fitz.open(path)
    result = []
    for i in range(src.page_count):
        out = fitz.open()
        out.insert_pdf(src, from_page=i, to_page=i)
        target = Path(output_dir) / f"page-{i+1}.pdf"
        out.save(target)
        out.close()
        result.append(target)
    src.close()
    return result


def extract_pages(path, pages, output):
    src = fitz.open(path)
    out = fitz.open()
    for n in pages:
        idx = n - 1
        if 0 <= idx < src.page_count:
            out.insert_pdf(src, from_page=idx, to_page=idx)
    out.save(output)
    out.close(); src.close()


def rotate_pdf(path, angle, output):
    src = fitz.open(path)
    for page in src:
        page.set_rotation((page.rotation + angle) % 360)
    src.save(output)
    src.close()


def compress_pdf(path, output):
    src = fitz.open(path)
    src.save(output, garbage=4, deflate=True, clean=True)
    src.close()


def pdf_to_images(path, output_dir, fmt="png"):
    out = []
    doc = fitz.open(path)
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=fitz.Matrix(1.7, 1.7), alpha=False)
        target = Path(output_dir) / f"page-{i+1}.{fmt.lower()}"
        pix.save(target)
        out.append(target)
    doc.close()
    return out


def images_to_pdf(paths, output):
    out = fitz.open()
    for p in paths:
        img = Image.open(p).convert("RGB")
        buf = io.BytesIO(); img.save(buf, format="PNG")
        dpi = img.info.get("dpi", (72, 72))[0] or 72
        page = out.new_page(width=img.width * 72 / dpi, height=img.height * 72 / dpi)
        page.insert_image(page.rect, stream=buf.getvalue())
    out.save(output)
    out.close()


def json_from_gemini(text):
    match = re.search(r"```(?:json)?\s*(.*?)\s*```", text, re.S)
    raw = match.group(1) if match else text.strip()
    return json.loads(raw)


def _set_page_size(section, width_pt, height_pt):
    section.page_width = Pt(width_pt)
    section.page_height = Pt(height_pt)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _hex_color(value, default="000000"):
    if not isinstance(value, str):
        return default
    value = value.strip().lstrip("#")
    if len(value) == 6 and all(c in "0123456789abcdefABCDEF" for c in value):
        return value.upper()
    return default


def _add_textbox(doc, left_pt, top_pt, width_pt, height_pt, element):
    """Add an editable, absolutely positioned Word text box.

    VML attributes such as id/style/type are plain VML attributes and must not
    be passed through python-docx qn(), which expects a prefix:tag name.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))

    run = p.add_run()
    pict = OxmlElement("w:pict")
    shape = etree.Element(f"{{{VML_NS}}}shape")
    shape.set("id", f"GeminiText_{id(element)}")
    shape.set("type", "#_x0000_t202")
    shape.set("style",
              f"position:absolute;margin-left:{left_pt}pt;margin-top:{top_pt}pt;"
              f"width:{max(width_pt, 1)}pt;height:{max(height_pt, 1)}pt;"
              "z-index:1;mso-wrap-style:none;mso-position-horizontal:absolute;"
              "mso-position-horizontal-relative:page;mso-position-vertical:absolute;"
              "mso-position-vertical-relative:page")
    shape.set("fillcolor", "white")
    shape.set("stroked", "f")
    textbox = etree.SubElement(shape, f"{{{VML_NS}}}textbox")
    textbox.set("inset", "0pt,0pt,0pt,0pt")
    txbx = OxmlElement("w:txbxContent")
    wp = OxmlElement("w:p")
    wr = OxmlElement("w:r")
    wt = OxmlElement("w:t")
    wt.text = str(element.get("text", ""))
    wr.append(wt); wp.append(wr); txbx.append(wp); textbox.append(txbx)
    shape.append(textbox); pict.append(shape); run._r.append(pict)

    rPr = OxmlElement("w:rPr")
    if element.get("bold"):
        rPr.append(OxmlElement("w:b"))
    if element.get("italic"):
        rPr.append(OxmlElement("w:i"))
    if element.get("underline"):
        rPr.append(OxmlElement("w:u"))
    color = OxmlElement("w:color"); color.set(qn("w:val"), _hex_color(element.get("color"))); rPr.append(color)
    sz = OxmlElement("w:sz")
    try: sz.set(qn("w:val"), str(max(2, int(float(element.get("font_size", 10.5)) * 2))))
    except (TypeError, ValueError): sz.set(qn("w:val"), "21")
    rPr.append(sz)
    wr.insert(0, rPr)
    return p


def _add_floating_picture(doc, image_path, left_pt, top_pt, width_pt, height_pt):
    """Insert an original PDF image as a floating Word picture.

    The PDF page is never rasterized as a screenshot.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Pt(max(width_pt, 1)))
    inline_xml = inline._inline

    doc_pr = inline_xml.find(qn("wp:docPr"))
    graphic = inline_xml.find(qn("a:graphic"))
    extent = inline_xml.find(qn("wp:extent"))
    if doc_pr is None or graphic is None or extent is None:
        return p

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0"); anchor.set("distB", "0")
    anchor.set("distL", "0"); anchor.set("distR", "0")
    anchor.set("simplePos", "0"); anchor.set("relativeHeight", "251658240")
    anchor.set("behindDoc", "0"); anchor.set("locked", "0")
    anchor.set("layoutInCell", "1"); anchor.set("allowOverlap", "1")

    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0"); simple.set("y", "0"); anchor.append(simple)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    off_h = OxmlElement("wp:posOffset"); off_h.text = str(int(left_pt * 12700))
    pos_h.append(off_h); anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    off_v = OxmlElement("wp:posOffset"); off_v.text = str(int(top_pt * 12700))
    pos_v.append(off_v); anchor.append(pos_v)

    anchor.append(extent)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(doc_pr)
    c_nv = inline_xml.find(qn("wp:cNvGraphicFramePr"))
    if c_nv is not None:
        anchor.append(c_nv)
    anchor.append(graphic)

    parent = inline_xml.getparent()
    if parent is not None:
        parent.replace(inline_xml, anchor)
    return p


def _extract_page_images(page, work_dir):
    items = []
    for idx, info in enumerate(page.get_images(full=True)):
        xref = info[0]
        try:
            data = page.parent.extract_image(xref)
            ext = data.get("ext", "png")
            path = Path(work_dir) / f"page-{page.number+1}-image-{idx}.{ext}"
            path.write_bytes(data["image"])
            rects = page.get_image_rects(xref)
            rect = rects[0] if rects else fitz.Rect(0, 0, 100, 100)
            items.append((idx, path, rect))
        except Exception:
            continue
    return items


def _render_table(doc, page, element, page_width, page_height):
    rows = element.get("rows") or []
    if not rows:
        return
    max_cols = max((len(r) for r in rows if isinstance(r, list)), default=0)
    if max_cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    y = float(element.get("y", 0) or 0) * page_height
    w = float(element.get("w", 1) or 1) * page_width
    try:
        for r, row in enumerate(rows):
            if not isinstance(row, list):
                row = [str(row)]
            for c in range(max_cols):
                value = row[c] if c < len(row) else ""
                cell = table.cell(r, c)
                cell.text = "" if value is None else str(value)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    if element.get("font_size"):
                        for run in p.runs:
                            run.font.size = Pt(float(element["font_size"]))
        if w > 0:
            table.autofit = False
            col_width = Inches(w / 72 / max_cols)
            for row in table.rows:
                for cell in row.cells:
                    cell.width = col_width
    except (TypeError, ValueError):
        pass
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(max(0, y))
    spacer.paragraph_format.space_after = Pt(0)


def gemini_layout_to_docx(source_pdf, layout, output):
    """Build an editable DOCX from Gemini's visual layout description.

    Gemini understands the PDF. PyMuPDF is used only for page geometry and original
    embedded image extraction. Whole PDF pages are never converted to images.
    """
    source_pdf = Path(source_pdf)
    output = Path(output)
    pages = layout.get("pages", []) if isinstance(layout, dict) else []
    if not isinstance(pages, list):
        pages = []
    pdf = fitz.open(source_pdf)
    doc = Document()
    try:
        if not pages:
            raise ValueError("Gemini returned no page layout")
        for index, page in enumerate(pdf):
            if index:
                doc.add_section(WD_SECTION.NEW_PAGE)
            section = doc.sections[-1]
            width_pt, height_pt = page.rect.width, page.rect.height
            _set_page_size(section, width_pt, height_pt)
            gem_page = pages[index] if index < len(pages) and isinstance(pages[index], dict) else {}
            image_items = _extract_page_images(page, output.parent)
            columns = gem_page.get("columns", [])
            elements = []
            if isinstance(columns, list):
                for col in columns:
                    if isinstance(col, dict) and isinstance(col.get("elements"), list):
                        elements.extend(e for e in col["elements"] if isinstance(e, dict))
            page_elements = gem_page.get("elements", [])
            if isinstance(page_elements, list):
                elements.extend(e for e in page_elements if isinstance(e, dict))
            elements.sort(key=lambda e: (
                float(e.get("y", 0) or 0),
                float(e.get("x", 0) or 0),
            ))
            for element in elements:
                kind = str(element.get("type", "text") or "text").lower()
                x = float(element.get("x", 0) or 0) * width_pt
                y = float(element.get("y", 0) or 0) * height_pt
                w = float(element.get("w", 0.9) or 0.9) * width_pt
                h = float(element.get("h", 0.04) or 0.04) * height_pt
                if kind == "text" and str(element.get("text", "") or "").strip():
                    _add_textbox(doc, x, y, w, h, element)
                elif kind == "image":
                    try:
                        idx = int(element.get("image_index", 0) or 0)
                    except (TypeError, ValueError):
                        idx = 0
                    matches = [item for item in image_items if item[0] == idx]
                    if matches:
                        _, image_path, _ = matches[0]
                        _add_floating_picture(doc, image_path, x, y, w, h)
                elif kind == "table":
                    _render_table(doc, page, element, width_pt, height_pt)
                elif kind == "line":
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(y)
                    p.paragraph_format.space_after = Pt(0)
                    pPr = p._p.get_or_add_pPr()
                    borders = OxmlElement("w:pBdr")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "6")
                    bottom.set(qn("w:space"), "0")
                    bottom.set(qn("w:color"), _hex_color(element.get("color", "808080")))
                    borders.append(bottom); pPr.append(borders)
        doc.save(output)
    finally:
        pdf.close()


def structured_to_docx(data, output):
    doc = Document()
    for block in data.get("blocks", []):
        kind = block.get("type", "paragraph")
        text = block.get("text", "")
        if kind == "heading":
            doc.add_heading(text, level=min(int(block.get("level", 1)), 9))
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    doc.save(output)


def structured_to_xlsx(data, output):
    wb = Workbook(); ws = wb.active; ws.title = "PDF"
    rows = data.get("rows", [])
    for r, row in enumerate(rows, 1):
        for c, value in enumerate(row, 1):
            ws.cell(r, c, value)
    wb.save(output)
