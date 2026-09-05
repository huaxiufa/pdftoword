from __future__ import annotations

import io
import json
import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import html
from PIL import Image

EMU_PER_INCH = 914400
PT_PER_INCH = 72


class CoordinateDocxRenderer:
    """Rebuild PDF pages as editable Word coordinate canvases."""

    def __init__(self, pdf_path: Path):
        self.pdf = fitz.open(pdf_path)

    def render_native(self, out_path: Path) -> None:
        doc = Document()
        for page_no, page in enumerate(self.pdf):
            self._start_page(doc, page, page_no)
            self._render_images(doc, page)
            self._render_native_text(doc, page)
        self.pdf.close()
        doc.save(out_path)

    def render_mixed(self, out_path: Path, ocr_pages: dict[int, Path]) -> None:
        """Render native PDF pages directly and OCR only scanned pages."""
        doc = Document()
        for page_no, page in enumerate(self.pdf):
            self._start_page(doc, page, page_no)
            if page_no in ocr_pages:
                self._render_images(doc, page)
                self._render_ocr(doc, page, ocr_pages[page_no])
            else:
                self._render_images(doc, page)
                self._render_native_text(doc, page)
        self.pdf.close()
        doc.save(out_path)

    def render(self, json_files: list[Path], out_path: Path) -> None:
        """Backward-compatible OCR-only renderer."""
        doc = Document()
        for page_no, json_path in enumerate(json_files):
            page_index = self._page_index(json_path, page_no)
            page = self.pdf[page_index]
            self._start_page(doc, page, page_no)
            self._render_images(doc, page)
            self._render_ocr(doc, page, json_path)
        self.pdf.close()
        doc.save(out_path)

    def _start_page(self, doc, page, page_no):
        section = doc.sections[0] if page_no == 0 else doc.add_section(WD_SECTION.NEW_PAGE)
        self._configure_section(section, page.rect.width, page.rect.height)

    def _configure_section(self, section, width_pt, height_pt):
        section.page_width = Pt(width_pt)
        section.page_height = Pt(height_pt)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(0)
        section.header_distance = section.footer_distance = Pt(0)

    def _render_ocr(self, doc, page, json_path):
        data = json.loads(json_path.read_text(encoding="utf-8"))
        if isinstance(data.get("res"), dict):
            data = data["res"]
        img_path = json_path.with_suffix(".png")
        with Image.open(img_path) as im:
            src_w, src_h = im.size
        sx, sy = page.rect.width / src_w, page.rect.height / src_h
        blocks = data.get("parsing_res_list") or []
        table_htmls = [
            item["pred_html"] for item in data.get("table_res_list") or []
            if isinstance(item, dict) and isinstance(item.get("pred_html"), str)
        ]
        table_index = 0
        for block in blocks:
            bbox = block.get("block_bbox") or block.get("bbox")
            if not bbox or len(bbox) < 4:
                continue
            bbox = [float(v) for v in bbox[:4]]
            bbox = [bbox[0] * sx, bbox[1] * sy, bbox[2] * sx, bbox[3] * sy]
            label = str(block.get("block_label") or block.get("sub_label") or "text").lower()
            if label in {"image", "figure", "chart"}:
                continue
            text = str(block.get("block_content") or "").strip()
            if label == "table":
                if "<table" in text.lower():
                    self._add_table_box(doc, text, bbox)
                elif table_index < len(table_htmls):
                    self._add_table_box(doc, table_htmls[table_index], bbox)
                    table_index += 1
                elif text:
                    self._add_text_box(doc, text, bbox)
            elif text:
                self._add_text_box(doc, text, bbox)

    def _render_images(self, doc, page):
        page_w, page_h = page.rect.width, page.rect.height
        seen = set()
        for img in page.get_images(full=True):
            xref = img[0]
            for rect in page.get_image_rects(xref):
                area = max(0, rect.width) * max(0, rect.height)
                if area >= page_w * page_h * 0.92:
                    continue
                key = (xref, round(rect.x0, 1), round(rect.y0, 1), round(rect.x1, 1), round(rect.y1, 1))
                if key in seen:
                    continue
                seen.add(key)
                try:
                    pix = fitz.Pixmap(self.pdf, xref)
                    if pix.alpha:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    self._add_floating_image(doc, pix.tobytes("png"), rect)
                except Exception:
                    pass

    def _render_native_text(self, doc, page):
        """Use the PDF's own text geometry; no OCR/model inference required."""
        blocks = page.get_text("dict").get("blocks", [])
        tables = []
        try:
            finder = page.find_tables()
            tables = list(finder.tables) if finder else []
        except Exception:
            tables = []

        table_rects = [fitz.Rect(t.bbox) for t in tables]
        for table in tables:
            try:
                rows = table.extract()
                if rows:
                    self._add_native_table(doc, rows, fitz.Rect(table.bbox))
            except Exception:
                pass

        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = str(span.get("text", ""))
                    bbox = span.get("bbox")
                    if not text.strip() or not bbox:
                        continue
                    rect = fitz.Rect(bbox)
                    if any(rect.intersects(trect) and rect.get_area() < trect.get_area() * 0.8 for trect in table_rects):
                        continue
                    self._add_text_box(
                        doc,
                        text,
                        bbox,
                        font_pt=float(span.get("size") or max(6, rect.height * 0.72)),
                        font_name=str(span.get("font") or "Arial"),
                    )

    def _add_native_table(self, doc, rows, rect):
        normalized = []
        for row in rows:
            normalized.append([(self._cell_text(v), False) for v in row])
        self._add_table_rows_box(doc, normalized, [rect.x0, rect.y0, rect.x1, rect.y1])

    @staticmethod
    def _cell_text(value):
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        return str(value).strip()

    def _add_text_box(self, doc, text, bbox, font_pt=None, font_name="Arial"):
        x0, y0, x1, y1 = bbox
        width = max(1, x1 - x0)
        height = max(1, y1 - y0)
        font_pt = font_pt or max(6, min(36, height * 0.72))
        p = doc.add_paragraph()
        self._zero_paragraph(p)
        pict = OxmlElement("w:pict")
        shape = self._shape(x0, y0, width, height)
        txbx = OxmlElement("v:textbox")
        content = OxmlElement("w:txbxContent")
        wp = OxmlElement("w:p")
        wr = OxmlElement("w:r")
        rp = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(max(2, int(font_pt * 2))))
        rp.append(sz)
        fonts = OxmlElement("w:rFonts")
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{key}"), font_name)
        rp.append(fonts)
        wr.append(rp)
        wt = OxmlElement("w:t")
        wt.text = text
        wr.append(wt)
        wp.append(wr)
        content.append(wp)
        txbx.append(content)
        shape.append(txbx)
        pict.append(shape)
        p._p.append(pict)

    def _add_table_box(self, doc, table_html, bbox):
        root = html.fromstring(table_html)
        table = root if root.tag.lower() == "table" else root.find(".//table")
        if table is None:
            return self._add_text_box(doc, root.text_content().strip(), bbox)
        rows = []
        for row in table.xpath(".//tr"):
            rows.append([(c.text_content().strip(), c.tag.lower() == "th") for c in row.xpath("./th|./td")])
        if rows:
            self._add_table_rows_box(doc, rows, bbox)

    def _add_table_rows_box(self, doc, rows, bbox):
        x0, y0, x1, y1 = bbox
        p = doc.add_paragraph()
        self._zero_paragraph(p)
        pict = OxmlElement("w:pict")
        shape = self._shape(x0, y0, x1 - x0, y1 - y0)
        txbx = OxmlElement("v:textbox")
        content = OxmlElement("w:txbxContent")
        tbl = OxmlElement("w:tbl")
        self._build_table(tbl, rows)
        content.append(tbl)
        txbx.append(content)
        shape.append(txbx)
        pict.append(shape)
        p._p.append(pict)

    def _build_table(self, tbl, rows):
        tbl_pr = OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "808080")
            borders.append(el)
        tbl_pr.append(borders)
        tbl.append(tbl_pr)
        for cells in rows:
            tr = OxmlElement("w:tr")
            for text, header in cells:
                tc = OxmlElement("w:tc")
                p = OxmlElement("w:p")
                r = OxmlElement("w:r")
                if header:
                    rp = OxmlElement("w:rPr")
                    rp.append(OxmlElement("w:b"))
                    r.append(rp)
                t = OxmlElement("w:t")
                t.text = text
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

    def _add_floating_image(self, doc, blob, rect):
        x0, y0, x1, y1 = rect.x0, rect.y0, rect.x1, rect.y1
        p = doc.add_paragraph()
        self._zero_paragraph(p)
        run = p.add_run()
        inline_shape = run.add_picture(
            io.BytesIO(blob), width=Pt(max(1, x1 - x0)), height=Pt(max(1, y1 - y0))
        )
        inline = inline_shape._inline
        anchor = OxmlElement("wp:anchor")
        for k, v in {"distT": "0", "distB": "0", "distL": "0", "distR": "0", "simplePos": "0", "relativeHeight": "0", "behindDoc": "0", "locked": "0", "layoutInCell": "1", "allowOverlap": "1"}.items():
            anchor.set(k, v)
        simple = OxmlElement("wp:simplePos")
        simple.set("x", "0")
        simple.set("y", "0")
        anchor.append(simple)
        for tag, value in (("positionH", x0), ("positionV", y0)):
            pos = OxmlElement(f"wp:{tag}")
            pos.set("relativeFrom", "page")
            off = OxmlElement("wp:posOffset")
            off.text = str(int(value * EMU_PER_INCH / PT_PER_INCH))
            pos.append(off)
            anchor.append(pos)
        extent = OxmlElement("wp:extent")
        extent.set("cx", str(int((x1 - x0) * EMU_PER_INCH / PT_PER_INCH)))
        extent.set("cy", str(int((y1 - y0) * EMU_PER_INCH / PT_PER_INCH)))
        anchor.append(extent)
        eff = OxmlElement("wp:effectExtent")
        for k in ("l", "t", "r", "b"):
            eff.set(k, "0")
        anchor.append(eff)
        anchor.append(OxmlElement("wp:wrapNone"))
        for child in list(inline):
            anchor.append(child)
        inline.getparent().replace(inline, anchor)

    def _shape(self, x, y, width, height):
        shape = OxmlElement("v:shape")
        shape.set(qn("id"), "_x0000_s" + str(abs(hash((x, y, width, height))) % 900000 + 100000))
        shape.set(qn("type"), "#_x0000_t202")
        shape.set("style", f"position:absolute;margin-left:{x}pt;margin-top:{y}pt;width:{width}pt;height:{height}pt;mso-wrap-style:none")
        shape.set(qn("stroked"), "f")
        shape.set(qn("filled"), "f")
        return shape

    def _zero_paragraph(self, p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

    def _page_index(self, path, fallback):
        m = re.search(r"page-(\d+)", path.name)
        return int(m.group(1)) if m else fallback
